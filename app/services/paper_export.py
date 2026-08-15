from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.schemas.question import Question, QuestionType


def export_paper_pdf(questions: list[Question], title: str, output_path: str) -> str:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    story = [Paragraph(title, styles["Title"]), Spacer(1, 0.3 * inch)]

    total_marks = sum(q.marks for q in questions)
    story.append(Paragraph(f"Total Marks: {total_marks}", styles["Normal"]))
    story.append(Spacer(1, 0.25 * inch))

    for i, q in enumerate(questions, start=1):
        story.append(Paragraph(f"{i}. {q.text} [{q.marks} marks]", styles["Normal"]))
        if q.question_type == QuestionType.MCQ and q.options:
            for j, option in enumerate(q.options):
                letter = chr(ord("a") + j)
                story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;({letter}) {option}", styles["Normal"]))
        story.append(Spacer(1, 0.15 * inch))

    doc.build(story)
    return output_path


def export_paper_docx(questions: list[Question], title: str, output_path: str) -> str:
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Total Marks: {sum(q.marks for q in questions)}")

    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q.text} [{q.marks} marks]")
        if q.question_type == QuestionType.MCQ and q.options:
            for j, option in enumerate(q.options):
                letter = chr(ord("a") + j)
                doc.add_paragraph(f"     ({letter}) {option}")

    doc.save(output_path)
    return output_path
